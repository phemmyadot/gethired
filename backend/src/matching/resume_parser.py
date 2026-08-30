"""
Resume parser: extract plain text from PDF and DOCX files.
"""
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_pdf(path: str) -> str:
    """Extract text from a PDF resume using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed for {path}: {e}")
        return ""


def extract_docx(path: str) -> str:
    """Extract text from a DOCX resume using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {path}: {e}")
        return ""


def extract_resume_text(path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_docx(path)
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        logger.warning(f"Unsupported resume format: {ext}")
        return ""


def clean_text(text: str) -> str:
    """Normalize whitespace and remove junk characters."""
    import re
    text = re.sub(r"\s{3,}", "\n\n", text)
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)  # ASCII printable + newlines
    return text.strip()
